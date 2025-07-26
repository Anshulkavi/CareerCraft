# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from .models import ResumeData
# from .utils import extract_email, extract_phone, extract_name, extract_skills, extract_experience
# from .job_finder import scrape_internshala_jobs
# from django.core.validators import EmailValidator
# from django.core.exceptions import ValidationError
# import PyPDF2
# import docx


# # ---------------- EMAIL VALIDATOR ----------------
# def validate_email(email):
#     validator = EmailValidator()
#     try:
#         validator(email)
#         return True
#     except ValidationError:
#         return False

# # ---------------- TEXT EXTRACTOR ----------------
# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return "\n".join(para.text for para in doc.paragraphs)

# # ---------------- SKILL MATCHER ----------------
# def count_matching_skills(job, skills):
#     fields = [
#         job.get("title", ""),
#         job.get("location", ""),
#         job.get("salary", "")
#     ]
#     combined = " ".join(fields).lower()
#     return sum(1 for skill in skills if skill.lower() in combined)

# # ---------------- RESUME UPLOAD VIEW ----------------
# @csrf_exempt
# def upload_resume(request):
#     if request.method == 'POST' and request.FILES.get('resume'):
#         try:
#             f = request.FILES['resume']
#             ext = f.name.split('.')[-1].lower()

#             if f.size > 10 * 1024 * 1024:
#                 return JsonResponse({'error': 'File too large. Maximum size is 10MB.'}, status=400)

#             text = ""
#             if ext == 'pdf':
#                 reader = PyPDF2.PdfReader(f)
#                 for p in reader.pages:
#                     text += p.extract_text() or ""
#             elif ext == 'docx':
#                 text = extract_text_from_docx(f)
#             else:
#                 return JsonResponse({'error': 'Only PDF or DOCX files are allowed.'}, status=400)

#             name = extract_name(text)
#             email = extract_email(text)
#             phone = extract_phone(text)
#             skills = extract_skills(text)
#             experience = extract_experience(text)

#             if not email or '@' not in email or '.' not in email:
#                 email = None 
#             if not skills:
#                 return JsonResponse({'error': 'No skills found in the resume.'}, status=400)

#             # Save extracted data
#             ResumeData.objects.create(
#                 name=name or "Not specified",
#                 email=email,
#                 phone=phone or "Not specified",
#                 skills=skills,
#                 experience=experience or "Not specified"
#             )

#             # Fetch jobs
#             all_jobs = []
#             for skill in skills:
#                 all_jobs += scrape_internshala_jobs([skill])

#             # Count matches
#             for job in all_jobs:
#                 job["matching_skills"] = count_matching_skills(job, skills)

#             unique_jobs = {(job['title'], job.get('company_name', 'Unknown')): job for job in all_jobs}
#             all_jobs = list(unique_jobs.values())
#             sorted_jobs = sorted(all_jobs, key=lambda x: -x.get("matching_skills", 0))
#             top_matches = sorted_jobs[:5]

#             return JsonResponse({
#                 'message': 'Resume processed successfully!',
#                 'matches': top_matches,
#                 'extracted': {
#                     'name': name or "Not specified",
#                     'email': email or "Not specified",
#                     'phone': phone or "Not specified",
#                     'skills': skills,
#                     'experience': experience or "Not specified"
#                 }
#             })

#         except Exception as e:
#             # Critical: Catch unhandled exceptions to return proper JSON
#             return JsonResponse({'error': f'Internal server error: {str(e)}'}, status=500)

#     return JsonResponse({'error': 'Invalid request. POST with resume file required.'}, status=400)

# @csrf_exempt
# def test_cors(request):
#     return JsonResponse({"status": "ok"})

# def health_check(request):
#     return JsonResponse({'status': 'ok'}, status=200)


# 2nd code use when 3rd doesnt work

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import ResumeData
from .utils import extract_name, extract_email, extract_phone, extract_skills, extract_experience
from .job_finder import fetch_relevant_jobs, count_matching_skills
import PyPDF2
import docx
from django.core.validators import EmailValidator
from django.core.exceptions import ValidationError
from .models import Resume
from .serializers import ResumeSerializer
from rest_framework.decorators import api_view
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView
from rest_framework.response import Response

# ---------------- EMAIL VALIDATOR ----------------
def validate_email(email):
    validator = EmailValidator()
    try:
        validator(email)
        return True
    except ValidationError:
        return False

# # ---------------- TEXT EXTRACTOR ----------------
def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join(para.text for para in doc.paragraphs)


@csrf_exempt
def upload_resume(request):
    if request.method == 'POST' and request.FILES.get('resume'):
        try:
            f = request.FILES['resume']
            ext = f.name.split('.')[-1].lower()

            if f.size > 10 * 1024 * 1024:
                return JsonResponse({'error': 'File too large. Maximum size is 10MB.'}, status=400)

            text = ""
            if ext == 'pdf':
                reader = PyPDF2.PdfReader(f)
                for p in reader.pages:
                    text += p.extract_text() or ""
            elif ext == 'docx':
                text = extract_text_from_docx(f)
            else:
                return JsonResponse({'error': 'Only PDF or DOCX files are allowed.'}, status=400)

            # Extract candidate info
            name = extract_name(text)
            email = extract_email(text)
            phone = extract_phone(text)
            skills = extract_skills(text)
            experience = extract_experience(text)

            # Basic email validation
            if not email or '@' not in email or '.' not in email:
                email = None

            if not skills:
                return JsonResponse({'error': 'No skills found in the resume.'}, status=400)

            # Save to DB
            ResumeData.objects.create(
                name=name or "Not specified",
                email=email,
                phone=phone or "Not specified",
                skills=skills,
                experience=experience or "Not specified"
            )

            # Prioritize these skills and get top 5
            priority_skills = ['python', 'django', 'react']
            top_skills = sorted(skills, key=lambda s: s.lower() not in priority_skills)[:5]

            # Fetch relevant jobs
            job_results = fetch_relevant_jobs(top_skills)

            # Calculate matching skill count per job
            for job in job_results:
                job["matching_skills"] = count_matching_skills(job, top_skills)

            # Deduplicate jobs by (title, company_name)
            unique_jobs = {(job['title'], job.get('company_name', 'Unknown')): job for job in job_results}
            sorted_jobs = sorted(unique_jobs.values(), key=lambda x: -x.get("matching_skills", 0))[:5]

            return JsonResponse({
                'message': 'Resume processed successfully!',
                'matches': sorted_jobs,
                'extracted': {
                    'name': name or "Not specified",
                    'email': email or "Not specified",
                    'phone': phone or "Not specified",
                    'skills': top_skills,
                    'experience': experience or "Not specified"
                }
            })

        except Exception as e:
            return JsonResponse({'error': f'Internal server error: {str(e)}'}, status=500)

    return JsonResponse({'error': 'Invalid request. POST with resume file required.'}, status=400)

class SaveResumeView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        user_id = str(request.user.id)
        title = request.data.get("title")
        data = request.data.get('data', {})  # ✅ default to empty dict

        if not title or not data:
            return Response({"error": "Missing title or data"}, status=400)

        # Overwrite if same user + title exists
        existing = Resume.objects(user_id=user_id, title=title).first()
        if existing:
            existing.data = data
            existing.save()
        else:
            Resume(user_id=user_id, title=title, data=data).save()

        return Response({"message": "Resume saved successfully"})


class GetResumeView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user_id = str(request.user.id)

        # Fetch all resumes for this user, sorted by latest updated
        resumes = Resume.objects(user_id=user_id).order_by("-updated_at")

        resume_list = [
            {
                "id": str(resume.id),
                "title": resume.title,
                "created_at": resume.created_at.strftime("%Y-%m-%d %H:%M"),
                "updated_at": resume.updated_at.strftime("%Y-%m-%d %H:%M"),
            }
            for resume in resumes
        ]

        return Response(resume_list)
    
@api_view(['GET'])
def get_resume_by_id(request, resume_id):
    try:
        resume = Resume.objects.get(id=resume_id)
        serializer = ResumeSerializer(resume)
        return Response(serializer.data, status=200)
    except Resume.DoesNotExist:
        return Response({"error": "Resume not found"}, status=404)


    
@csrf_exempt
def test_cors(request):
    return JsonResponse({"status": "ok"})

def health_check(request):
    return JsonResponse({'status': 'ok'}, status=200)